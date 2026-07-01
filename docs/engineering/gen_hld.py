#!/usr/bin/env python3
"""
Generate Rootline_HLD_v1.docx — High-Level Design document.
Run: python3 gen_hld.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette ────────────────────────────────────────────────────────────────────
G_DARK  = RGBColor(0x0D, 0x4A, 0x28)
G_MID   = RGBColor(0x1A, 0x6B, 0x3C)
G_LIGHT = RGBColor(0xE8, 0xF4, 0xED)
GOLD    = RGBColor(0x9A, 0x6B, 0x00)
GREY    = RGBColor(0x5A, 0x5A, 0x5A)
GREY_LT = RGBColor(0xF4, 0xF4, 0xF4)
BLACK   = RGBColor(0x1C, 0x1C, 0x1C)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
AWS_ORG = RGBColor(0xFF, 0x99, 0x00)   # AWS orange
BLUE    = RGBColor(0x00, 0x52, 0x9B)   # tech blue for URLs/references
RED_LT  = RGBColor(0xB2, 0x22, 0x22)

# ── XML helpers ────────────────────────────────────────────────────────────────
def sp(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    spc = OxmlElement('w:spacing')
    spc.set(qn('w:before'), str(before))
    spc.set(qn('w:after'),  str(after))
    pPr.append(spc)

def ind(para, left=0, right=0, hanging=0):
    pPr = para._p.get_or_add_pPr()
    el  = OxmlElement('w:ind')
    if left:    el.set(qn('w:left'),    str(left))
    if right:   el.set(qn('w:right'),   str(right))
    if hanging: el.set(qn('w:hanging'), str(hanging))
    pPr.append(el)

def border_bottom(para, color='1A6B3C', size=6):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), str(size))
    b.set(qn('w:space'), '4');    b.set(qn('w:color'), color)
    pBdr.append(b); pPr.append(pBdr)

def border_left(para, color='9A6B00', size=18):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    l = OxmlElement('w:left')
    l.set(qn('w:val'), 'single'); l.set(qn('w:sz'), str(size))
    l.set(qn('w:space'), '12');   l.set(qn('w:color'), color)
    pBdr.append(l); pPr.append(pBdr)

def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

def set_cell_border(cell, color='CCCCCC'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4')
        el.set(qn('w:space'),'0');    el.set(qn('w:color'), color)
        tcB.append(el)
    tcPr.append(tcB)

def no_border_cell(cell):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'none'); el.set(qn('w:sz'),'0')
        el.set(qn('w:space'),'0'); el.set(qn('w:color'),'FFFFFF')
        tcB.append(el)
    tcPr.append(tcB)

# ── Document setup ─────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width    = Inches(8.5); sec.page_height   = Inches(11)
sec.left_margin   = Inches(1);   sec.right_margin  = Inches(1)
sec.top_margin    = Inches(0.9); sec.bottom_margin = Inches(0.9)
doc.styles['Normal'].font.name = 'Georgia'
doc.styles['Normal'].font.size = Pt(11)
doc.styles['Normal'].font.color.rgb = BLACK

# ── Helper functions ───────────────────────────────────────────────────────────
def cover_title(text, size=36, color=None):
    p = doc.add_paragraph(); sp(p, before=0, after=10)
    r = p.add_run(text)
    r.font.name='Arial'; r.font.size=Pt(size)
    r.font.bold=True; r.font.color.rgb = color or G_MID
    return p

def section_h1(num, title):
    p = doc.add_paragraph()
    p._p.get_or_add_pPr().set(qn('w:pageBreakBefore'), '1')
    sp(p, before=0, after=100); border_bottom(p, '1A6B3C', 10)
    r1 = p.add_run(f'{num}.  ')
    r1.font.name='Arial'; r1.font.size=Pt(20); r1.font.bold=True; r1.font.color.rgb=G_DARK
    r2 = p.add_run(title)
    r2.font.name='Arial'; r2.font.size=Pt(20); r2.font.bold=True; r2.font.color.rgb=G_MID
    return p

def h2(text, color=None):
    p = doc.add_paragraph(); sp(p, before=220, after=70)
    r = p.add_run(text)
    r.font.name='Arial'; r.font.size=Pt(13); r.font.bold=True
    r.font.color.rgb = color or G_DARK
    return p

def h3(text):
    p = doc.add_paragraph(); sp(p, before=160, after=50)
    r = p.add_run(text)
    r.font.name='Arial'; r.font.size=Pt(11); r.font.bold=True; r.font.italic=True
    r.font.color.rgb = G_MID
    return p

def body(text, before=50, after=90, italic=False, bold=False, color=None, size=11):
    p = doc.add_paragraph(); sp(p, before=before, after=after)
    r = p.add_run(text)
    r.font.name='Georgia'; r.font.size=Pt(size)
    r.font.italic=italic; r.font.bold=bold
    r.font.color.rgb = color or BLACK
    return p

def note(text):
    p = doc.add_paragraph(); sp(p, before=100, after=100)
    border_left(p, '9A6B00', 18); ind(p, left=420)
    r = p.add_run(text)
    r.font.name='Georgia'; r.font.size=Pt(10); r.font.italic=True
    r.font.color.rgb = GREY
    return p

def eyebrow(text):
    p = doc.add_paragraph(); sp(p, before=200, after=40)
    r = p.add_run(text.upper())
    r.font.name='Arial'; r.font.size=Pt(8); r.font.bold=True
    r.font.all_caps=True; r.font.color.rgb=G_MID
    return p

def bullet(text, level=0, color=None):
    p = doc.add_paragraph(); sp(p, before=30, after=30)
    left = str(360 + level*360)
    ind(p, left=int(left), hanging=240)
    rb = p.add_run('•  ')
    rb.font.name='Arial'; rb.font.size=Pt(11)
    rb.font.color.rgb = G_MID; rb.font.bold=True
    rt = p.add_run(text)
    rt.font.name='Georgia'; rt.font.size=Pt(11)
    rt.font.color.rgb = color or BLACK
    return p

def numbered(num, text):
    p = doc.add_paragraph(); sp(p, before=40, after=40)
    ind(p, left=480, hanging=300)
    rn = p.add_run(f'{num}.  ')
    rn.font.name='Arial'; rn.font.size=Pt(11); rn.font.bold=True; rn.font.color.rgb=G_MID
    rt = p.add_run(text)
    rt.font.name='Georgia'; rt.font.size=Pt(11); rt.font.color.rgb=BLACK
    return p

def thin_rule(before=80, after=80):
    p = doc.add_paragraph(); sp(p, before=before, after=after)
    border_bottom(p, 'DDDDDD', 4); p.add_run('')

def code_block(text):
    p = doc.add_paragraph(); sp(p, before=60, after=60)
    ind(p, left=360, right=360)
    r = p.add_run(text)
    r.font.name='Courier New'; r.font.size=Pt(9); r.font.color.rgb=G_DARK
    return p

def mk_table(headers, rows, header_bg='0D4A28', alt_bg='F7FBF9'):
    cols = len(headers)
    t = doc.add_table(rows=1+len(rows), cols=cols)
    t.style='Table Grid'
    # header
    for j,h in enumerate(headers):
        cell = t.rows[0].cells[j]
        set_cell_bg(cell, header_bg); set_cell_border(cell,'CCCCCC')
        pp = cell.paragraphs[0]; sp(pp,before=80,after=80); ind(pp,left=100)
        r = pp.add_run(h)
        r.font.name='Arial'; r.font.size=Pt(10); r.font.bold=True; r.font.color.rgb=WHITE
    # rows
    for i,row in enumerate(rows):
        for j,val in enumerate(row):
            cell = t.rows[i+1].cells[j]
            bg = 'FFFFFF' if i%2==0 else alt_bg
            set_cell_bg(cell, bg); set_cell_border(cell,'DDDDDD')
            pp = cell.paragraphs[0]; sp(pp,before=70,after=70); ind(pp,left=100)
            # bold first column
            r = pp.add_run(str(val))
            r.font.name = 'Arial' if j==0 else 'Georgia'
            r.font.size = Pt(10)
            r.font.bold = (j==0)
            r.font.color.rgb = G_DARK if j==0 else BLACK
    doc.add_paragraph()
    return t

def aws_badge(service, detail=''):
    """Inline AWS service callout line."""
    p = doc.add_paragraph(); sp(p, before=25, after=25); ind(p, left=360)
    ra = p.add_run('AWS  ')
    ra.font.name='Arial'; ra.font.size=Pt(9); ra.font.bold=True; ra.font.color.rgb=AWS_ORG
    rs = p.add_run(service)
    rs.font.name='Arial'; rs.font.size=Pt(10); rs.font.bold=True; rs.font.color.rgb=BLACK
    if detail:
        rd = p.add_run(f'  —  {detail}')
        rd.font.name='Georgia'; rd.font.size=Pt(10); rd.font.color.rgb=GREY
    return p

# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
cover_title('ROOTLINE', 40, G_MID)
cover_title('High-Level Design — v1', 22, BLACK)

sep = doc.add_paragraph(); sp(sep, before=10, after=60)
border_bottom(sep, '1A6B3C', 8); sep.add_run('')

meta = doc.add_paragraph(); sp(meta, before=0, after=20)
r = meta.add_run('Document Type: ')
r.font.name='Arial'; r.font.size=Pt(10); r.font.bold=True; r.font.color.rgb=GREY
r2 = meta.add_run('High-Level Engineering Design')
r2.font.name='Arial'; r2.font.size=Pt(10); r2.font.color.rgb=BLACK

for label, val in [
    ('AWS Account:', '741705262274  (us-east-1 primary)'),
    ('Bedrock Region:', 'us-east-1'),
    ('Status:', 'Draft — v1.0 — June 2026'),
    ('Author:', 'Chief Technology Officer, Rootline'),
]:
    p = doc.add_paragraph(); sp(p, before=0, after=12)
    r = p.add_run(label + '  ')
    r.font.name='Arial'; r.font.size=Pt(10); r.font.bold=True; r.font.color.rgb=GREY
    r2 = p.add_run(val)
    r2.font.name='Arial'; r2.font.size=Pt(10); r2.font.color.rgb=BLACK

doc.add_paragraph()
body(
    'This document describes the high-level architecture for Rootline v1. '
    'All infrastructure runs on AWS. The AI layer is built on Amazon Bedrock and is '
    'designed as an event-driven multi-agent system that supports the product from '
    'MVP through to v2 and beyond without re-architecture.',
    before=80, after=0, italic=True, color=GREY, size=10
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — ARCHITECTURE PRINCIPLES
# ══════════════════════════════════════════════════════════════════════════════
section_h1('1', 'Architecture Principles')
body(
    'Every engineering decision in v1 is governed by six principles, each mapping directly '
    'to a product design constraint established in the Working Backward document.'
)

principles = [
    ('Serverless-first',
     'Lambda + managed services over EC2. No servers to patch, near-zero idle cost, '
     'automatic scaling from 1 to 10,000 concurrent requests without configuration.'),
    ('Event-driven throughout',
     'All cross-service communication goes through Amazon EventBridge. Services emit events '
     'and react to events. No direct service-to-service synchronous coupling beyond the API layer. '
     'This is the foundation of the agentic AI architecture.'),
    ('Offline-first for suppliers',
     'The supplier portal is a Progressive Web App (PWA) with a Service Worker that queues '
     'submissions in IndexedDB when offline and syncs automatically on reconnect. '
     'This is non-negotiable for farm-based supply chain participants.'),
    ('Agents over hard-coded rules',
     'Business logic that involves judgment — fraud signals, deviation severity, '
     'notification routing — is handled by Amazon Bedrock agents, not if/else code. '
     'This means the intelligence can be improved by upgrading the model or adding tools, '
     'not by redeploying application code.'),
    ('One-way doors are locked in v1',
     'The retailer-canonical lot identity model, the SHA-256 audit log, the static KDE '
     'data schema, and the white-label brand config structure are fixed. '
     'These cannot be changed after the first retailer onboards without breaking data continuity.'),
    ('Scale to v2 without re-architecture',
     'Every v2 capability — IoT sensor ingestion, AI photo verification, supplier risk scoring, '
     'blockchain anchoring — is designed as an additive agent or event consumer. '
     'The v1 event schema and data model are designed to support these without migration.'),
]
for i,(title,desc) in enumerate(principles, 1):
    numbered(i, f'{title}:  {desc}')

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — SYSTEM OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
section_h1('2', 'System Overview')
body(
    'Rootline v1 is composed of five architectural layers. Each layer has a single, '
    'well-defined responsibility. Layers communicate downward synchronously (client → API → service → data) '
    'and laterally asynchronously (service → EventBridge → agent → service).'
)

code_block(
'''┌──────────────────────────────────────────────────────────────────────┐
│  LAYER 1 — CLIENT                                                    │
│  Supplier PWA  │  Retailer Dashboard  │  Consumer QR Page            │
│  (Next.js)         (Next.js)              (Static / CDN)             │
├──────────────────────────────────────────────────────────────────────┤
│  LAYER 2 — API                                                        │
│  Amazon API Gateway  (REST + WebSocket)  +  AWS WAF                  │
├──────────────────────────────────────────────────────────────────────┤
│  LAYER 3 — SERVICES  (AWS Lambda — stateless microservices)          │
│  Auth │ Retailer │ Product │ Lot │ Handoff │ QR │ Export │ Notify    │
├──────────────────────────────────────────────────────────────────────┤
│  LAYER 4 — AI AGENTS  (Amazon Bedrock + Step Functions + SQS)        │
│  Verification │ Cold Chain Alert │ Compliance │ Notification Orch.   │
├──────────────────────────────────────────────────────────────────────┤
│  LAYER 5 — DATA                                                       │
│  RDS PostgreSQL │ DynamoDB (audit) │ S3 (media) │ ElastiCache (Redis)│
└──────────────────────────────────────────────────────────────────────┘

                        EVENT BUS
               Amazon EventBridge (runs across L3 ↔ L4)
'''
)

body(
    'The EventBridge event bus is the nervous system of the platform. Every significant domain '
    'action — HandoffSubmitted, TemperatureDeviation, ChainComplete, ExportRequested — is published '
    'as a structured event. AI agents subscribe to relevant event patterns and act autonomously. '
    'This design means new capabilities (v2 agents) can be added by subscribing to existing events '
    'without touching v1 service code.'
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — CLIENT LAYER
# ══════════════════════════════════════════════════════════════════════════════
section_h1('3', 'Client Layer')
body(
    'Three distinct web surfaces serve the three personas. All are built in Next.js and '
    'served through Amazon CloudFront. No native mobile apps are built in v1.'
)

h2('3.1  Supplier Mobile Portal  (go.rootline.io)')
aws_badge('App Runner', 'hosts Next.js server-side rendered application')
aws_badge('CloudFront', 'CDN layer — <3s load on 3G globally')
body(
    'A Progressive Web App (PWA) optimised for mid-range Android on 3G. '
    'The Service Worker caches the application shell and task queue UI so the app is fully '
    'functional with no internet connection. Submitted handoff data is stored in '
    'the browser\'s IndexedDB and synced to the API when connectivity is restored.',
    before=60
)
bullet('Framework: Next.js 15 (App Router)')
bullet('Auth: SMS OTP → short-lived JWT stored in HttpOnly cookie')
bullet('Offline: PWA Service Worker + IndexedDB submission queue')
bullet('Photo capture: native browser camera API (getUserMedia) — no native app required')
bullet('GPS: navigator.geolocation + EXIF extraction from photo metadata')
bullet('OCR: camera capture → sends image to Textract Lambda on submit (cert number extraction)')
bullet('Max page load target: 3 seconds on 3G  |  Submit flow: under 60 seconds total')

h2('3.2  Retailer Web Dashboard  (app.rootline.io)')
aws_badge('App Runner', 'Next.js SSR — server-side rendered for fast first paint')
aws_badge('API Gateway WebSocket', 'real-time lot status updates without polling')
body(
    'Server-side rendered dashboard giving the CPO and store managers full operational '
    'visibility. Real-time lot status updates arrive via a persistent WebSocket connection '
    '— the dashboard reflects a cold chain alert within 5 seconds of a supplier submitting '
    'a deviating temperature reading.',
    before=60
)
bullet('Framework: Next.js 15 (App Router, server components)')
bullet('Real-time: API Gateway WebSocket API → Lambda → pushes status events to connected clients')
bullet('Auth: Amazon Cognito User Pool (email + password, MFA optional)')
bullet('Roles: Root Admin (full org) → Store Manager (store-scoped) — enforced at API layer')
bullet('FSMA export: triggers Compliance Agent via API → polling or WebSocket for completion signal')

h2('3.3  Consumer QR Experience  (verify.rootline.io / custom retailer domain)')
aws_badge('S3 + CloudFront', 'static HTML pages generated per lot, served from edge')
aws_badge('CloudFront Functions', 'URL routing and white-label brand injection at edge')
body(
    'When a retailer generates a QR code, Rootline pre-renders a static HTML page for '
    'that lot and uploads it to S3. CloudFront serves it from the nearest edge location '
    'globally. Page load is under 1 second. The page is white-labeled — all brand tokens '
    '(logo URL, primary colour, store name) are injected at generation time and baked into '
    'the static HTML. No runtime branding logic runs on the consumer request path.',
    before=60
)
bullet('Static generation: QR Service Lambda renders the HTML template with lot data + retailer brand at QR creation time')
bullet('Hosting: S3 bucket + CloudFront distribution with edge caching (TTL: 7 days, invalidated if lot data updates)')
bullet('Custom domains: Retailers can CNAME their own domain (e.g., verify.greenleafgrocers.com) to the CloudFront distribution')
bullet('No login required — public page, unique URL per lot')
bullet('Shareability: URL is stable and shareable via iOS/Android share sheet')

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — API LAYER
# ══════════════════════════════════════════════════════════════════════════════
section_h1('4', 'API Layer')
aws_badge('API Gateway', 'REST API (v1) + WebSocket API (real-time) — regional endpoint')
aws_badge('AWS WAF', 'rate limiting, IP blocking, SQL injection, XSS protection')
aws_badge('Route 53', 'DNS — api.rootline.io → API Gateway custom domain')

body(
    'All client-server communication goes through Amazon API Gateway. '
    'The REST API handles all synchronous requests. The WebSocket API manages persistent '
    'connections for the retailer dashboard real-time feed. WAF sits in front of both.',
    before=60
)

h2('4.1  REST API — Top-Level Resource Groups')
mk_table(
    ['Resource', 'Methods', 'Auth', 'Description'],
    [
        ['/auth/otp',         'POST',          'None',              'Request SMS OTP for a phone number'],
        ['/auth/verify',      'POST',          'None',              'Verify OTP → return JWT'],
        ['/auth/session',     'POST',          'Cognito',           'Retailer login → Cognito JWT'],
        ['/retailers',        'GET / PATCH',   'Cognito (Admin)',   'Brand config, subscription tier'],
        ['/stores',           'GET / POST',    'Cognito',           'Store management'],
        ['/products',         'GET / POST / PATCH', 'Cognito',      'Product/SKU catalog'],
        ['/lots',             'GET / POST',    'Cognito',           'Supply chain run management'],
        ['/lots/{id}/events', 'GET',           'Cognito',           'Full event timeline for a lot'],
        ['/lots/{id}/qr',     'POST',          'Cognito',           'Generate consumer QR + static page'],
        ['/lots/{id}/export', 'POST',          'Cognito',           'Trigger FSMA 204 export (async)'],
        ['/handoffs',         'POST',          'Supplier JWT',      'Submit a handoff event (supplier)'],
        ['/handoffs/{id}',    'GET',           'Supplier JWT',      'Get task detail for a participant'],
        ['/consumer/{lotId}', 'GET',           'None (public)',     'Consumer QR page data (cached)'],
        ['/ws/dashboard',     'WebSocket',     'Cognito',           'Real-time lot status stream'],
    ]
)

h2('4.2  Authentication Model')
body(
    'Two separate authentication flows serve two distinct user types with different technical realities.',
    before=0
)

h3('Retailer Auth — Amazon Cognito')
bullet('Email + password registration, verified email required')
bullet('Cognito User Pool issues JWT (access + refresh tokens) on successful login')
bullet('MFA via TOTP optional for v1, enforced in v2')
bullet('JWT passed as Bearer token in Authorization header on all retailer API calls')
bullet('Role claims (root_admin, store_manager, store_id) embedded in JWT and enforced at Lambda authorizer')

h3('Supplier Auth — SMS OTP + Custom JWT')
bullet('No email or password required — phone number is the only identifier')
bullet('POST /auth/otp sends a 6-digit code via Amazon Pinpoint SMS (expires in 10 minutes)')
bullet('POST /auth/verify validates the OTP and issues a signed JWT (HS256, 24-hour TTL)')
bullet('JWT contains: participant_id, phone_hash, issued_at — no sensitive data in payload')
bullet('Lambda authorizer validates supplier JWT on all /handoffs routes')

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — SERVICE LAYER
# ══════════════════════════════════════════════════════════════════════════════
section_h1('5', 'Service Layer — Lambda Microservices')
aws_badge('AWS Lambda', 'all services are stateless Lambda functions (Node.js 22 runtime)')
aws_badge('Lambda Layers', 'shared utilities — DB client, JWT validation, event publisher')
aws_badge('AWS Secrets Manager', 'database credentials, API keys — never in environment variables')

body(
    'Each service is a single Lambda function (or a small set of functions) with a '
    'single well-defined domain responsibility. Services communicate with the data layer '
    'directly and emit domain events to EventBridge. They do not call each other directly.',
    before=60
)

services = [
    ('Auth Service',
     'OTP generation and verification, JWT issuance for suppliers, Cognito user management for retailers. '
     'Rate-limited: max 3 OTP requests per phone number per 15 minutes (enforced via ElastiCache counter).'),
    ('Retailer Service',
     'Retailer account management, brand configuration (logo, colour, display name), '
     'store creation and management, user/role management. '
     'Brand config is cached in ElastiCache — invalidated on every PATCH /retailers call.'),
    ('Product Service',
     'Product/SKU catalog management per retailer. Stores product name, category, '
     'expected temperature range (used as alert threshold by the Cold Chain Alert Agent), '
     'and optional product description.'),
    ('Lot Service',
     'Supply chain run (lot) creation with retailer-canonical lot code generation (RL-YYYY-NNNN). '
     'Participant assignment and sequencing. Publishes LotCreated event → triggers Notification Agent '
     'to send SMS invitations to all participants.'),
    ('Handoff Service',
     'Receives supplier handoff submissions. Validates required fields. '
     'Triggers async photo upload to S3. Computes SHA-256 hash of the event record '
     'and appends to the DynamoDB audit log (append-only). '
     'Publishes HandoffSubmitted event → triggers Verification Agent and Alert Agent.'),
    ('QR Service',
     'Generates consumer QR code and pre-renders the white-labeled static HTML page. '
     'Fetches lot event data + retailer brand config, renders HTML template, '
     'uploads to S3 as verify.rootline.io/{lot_id}/index.html. '
     'Generates QR code pointing to this URL. Returns print-ready label PDF to retailer.'),
    ('Export Service',
     'Triggered by retailer FSMA export request. Publishes ExportRequested event. '
     'Compliance Agent handles the async generation. Export Service polls for completion '
     'and returns a presigned S3 URL when ready (or pushes via WebSocket).'),
    ('Notification Service',
     'Thin dispatch layer that sends SMS via Amazon Pinpoint. '
     'Called by AI agents — not directly by other services. '
     'Records every sent notification in notification_log table for deduplication and audit.'),
]

for svc, desc in services:
    h3(svc)
    body(desc, before=20, after=80)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — AI AGENT LAYER
# ══════════════════════════════════════════════════════════════════════════════
section_h1('6', 'AI Agent Layer — Amazon Bedrock')
body(
    'The AI agent layer is the core differentiator of the Rootline architecture. '
    'Instead of hard-coding business logic (if temperature > threshold: send alert), '
    'each domain capability is implemented as an Amazon Bedrock agent — a model with '
    'a set of callable tools — orchestrated by AWS Step Functions and triggered by '
    'EventBridge events. This means intelligence improves by updating the model or '
    'adding tools, not by rewriting application code. New v2 agents subscribe to '
    'existing events without touching v1 code.'
)

note(
    'All agents run on Amazon Bedrock in us-east-1. '
    'Claude Haiku 4.5 (fast, low cost) handles high-volume reactive tasks. '
    'Claude Sonnet 4.6 handles reasoning-heavy tasks (compliance, fraud signals). '
    'Model selection per agent is configurable without code changes.'
)

# ── Agent Architecture Diagram ────────────────────────────────────────────────
h2('6.1  Agent Architecture')
code_block(
'''DOMAIN EVENTS (EventBridge)
│
├── HandoffSubmitted ──────────► VERIFICATION AGENT (Haiku 4.5)
│                                   Tools: verify_usda_cert | validate_gps
│                                          compute_sha256 | flag_anomaly
│                                   Output: VerificationComplete event
│
├── HandoffSubmitted ──────────► COLD CHAIN ALERT AGENT (Haiku 4.5)
│   (temperature present)           Tools: get_threshold | calc_deviation
│                                          update_lot_status | send_sms
│                                   Output: AlertSent | LotStatusUpdated
│
├── ExportRequested ───────────► COMPLIANCE AGENT (Sonnet 4.6)
│                      Step Functions orchestration
│                                   Tools: get_all_ctes | map_to_kde
│                                          generate_pdf | upload_s3 | notify
│                                   Output: ExportReady event + S3 URL
│
├── LotCreated / HandoffSubmitted → NOTIFICATION ORCHESTRATOR (Haiku 4.5)
│   ChainComplete / Overdue          Tools: get_overdue_participants
│   (EventBridge Scheduler)                 check_notification_history
│                                           send_sms | escalate
│                                   Output: NotificationSent events
│
└── [v2] HandoffSubmitted ─────► FRAUD DETECTION AGENT (Sonnet 4.6)
                                    Tools: analyze_photo_metadata
                                           detect_duplicate_image
                                           score_supplier_history
                                           flag_for_review
'''
)

# ── Agent Definitions ─────────────────────────────────────────────────────────
h2('6.2  Verification Agent')
aws_badge('Bedrock', 'Claude Haiku 4.5  —  high volume, low latency')
aws_badge('SQS', 'HandoffSubmitted events queued for reliable delivery')
body(
    'Triggered immediately on every HandoffSubmitted event. Runs three parallel verification '
    'checks and writes the result to the audit log. Does not block the supplier\'s submission — '
    'verification is asynchronous.',
    before=60
)
body('Tools available to this agent:', before=40, after=20, bold=True)
bullet('verify_usda_cert(cert_number, claimed_farm_name)  →  calls USDA Organic Integrity Database API; returns match/no-match/pending')
bullet('validate_gps_plausibility(lat, lon, participant_id)  →  checks submitted GPS against participant\'s registered address; flags if >50 miles')
bullet('compute_sha256(event_payload)  →  computes tamper-evident hash; writes to DynamoDB audit_log (append-only)')
bullet('flag_anomaly(lot_id, event_id, severity, reason)  →  creates alert record; publishes AnomalyDetected event')
body('Decision logic the agent applies:', before=40, after=20, bold=True)
bullet('Cert number not found in USDA registry: flag as AMBER — cert may be recently issued; retailer notified to verify directly')
bullet('Cert number found but farm name does not match registry: flag as RED — potential fraud; retailer alerted immediately')
bullet('GPS >50 miles from registered address: flag as AMBER — participant may have moved or used wrong device')
bullet('All checks pass: VerificationComplete event with status=VERIFIED published to EventBridge')

h2('6.3  Cold Chain Alert Agent')
aws_badge('Bedrock', 'Claude Haiku 4.5  —  sub-second response required')
aws_badge('EventBridge', 'triggered on HandoffSubmitted where temp_reading is present')
body(
    'Evaluates every temperature reading against the product\'s configured threshold. '
    'Classifies the deviation by severity. Updates lot status on the dashboard in real time '
    'and sends an SMS to the retailer within 30 seconds of the supplier\'s submission.',
    before=60
)
body('Severity classification logic:', before=40, after=20, bold=True)
mk_table(
    ['Deviation from threshold', 'Severity', 'Retailer action triggered'],
    [
        ['≤0°F (within spec)',          'GREEN',    'No alert — lot status remains green'],
        ['1–3°F above threshold',        'AMBER',    'SMS alert to retailer — "brief deviation detected, reviewing"'],
        ['4–8°F above threshold',        'AMBER',    'SMS alert + dashboard flag — retailer prompted to contact carrier'],
        ['>8°F above threshold',         'RED',      'Urgent SMS + dashboard RED — retailer prompted to consider refusing lot'],
        ['Reading implausibly flat',     'AMBER',    'Potential data integrity issue — agent flags for manual review'],
    ]
)

h2('6.4  Compliance Agent')
aws_badge('Bedrock', 'Claude Sonnet 4.6  —  reasoning-heavy document generation')
aws_badge('Step Functions', 'orchestrates multi-step export workflow with retries')
aws_badge('Lambda', 'PDF generation using headless Chromium (Lambda layer)')
body(
    'Triggered when a retailer clicks "Export FSMA 204 Report." '
    'Step Functions orchestrates the multi-step workflow with retry logic at each step. '
    'The agent structures all Critical Tracking Events as FSMA Key Data Elements '
    'and generates a human-readable PDF and a machine-readable CSV.',
    before=60
)
body('Workflow steps (Step Functions state machine):', before=40, after=20, bold=True)
numbered(1, 'Fetch all HandoffEvents for the requested lot from RDS')
numbered(2, 'Map each event to FSMA KDE fields (lot code, dates, GPS, supplier/recipient identifiers)')
numbered(3, 'Identify and flag any missing CTEs in the chain')
numbered(4, 'Generate PDF report using headless Chromium Lambda (retailer-branded cover page + KDE tables)')
numbered(5, 'Generate structured CSV with FDA-compliant field names and formats')
numbered(6, 'Upload both files to S3 with a 7-day presigned URL')
numbered(7, 'Publish ExportReady event → Notification Service sends download link to retailer via SMS/WebSocket')

h2('6.5  Notification Orchestrator Agent')
aws_badge('Bedrock', 'Claude Haiku 4.5')
aws_badge('EventBridge Scheduler', 'runs every hour to check for overdue handoffs')
aws_badge('Pinpoint', 'SMS delivery — primary notification channel for all parties')
body(
    'Manages the full notification lifecycle: initial SMS invitations, 24-hour reminders, '
    '48-hour retailer alerts, and chain completion notifications. '
    'The agent checks notification history before sending to prevent duplicate messages.',
    before=60
)
body('Notification waterfall the agent manages:', before=40, after=20, bold=True)
mk_table(
    ['Trigger', 'Recipient', 'Channel', 'Message'],
    [
        ['Lot created',              'All participants',  'SMS', 'Invitation with registration link'],
        ['Prior participant submits','Next participant',  'SMS', 'Your turn — lot is heading your way'],
        ['24h since trigger, no log','Participant + CPO', 'SMS', 'Reminder + escalation notice to retailer'],
        ['48h since trigger, no log','CPO only',         'SMS', 'ALERT — chain at risk, participant overdue'],
        ['72h since trigger, no log','CPO only',         'SMS', 'CRITICAL — chain broken, action required'],
        ['Chain complete',           'CPO',              'SMS', 'Chain complete — QR ready to generate'],
    ]
)
body('Agent deduplication logic:', before=40, after=20, bold=True)
bullet('Before every send, agent calls check_notification_history(participant_id, lot_id, notification_type)')
bullet('If the same notification type was sent in the last 23 hours, send is skipped')
bullet('All sent notifications logged in notification_log table regardless of outcome')

h2('6.6  Fraud Detection Agent  (v2 — designed now, activated later)')
note(
    'This agent is not live in v1. The event stream, photo storage, and supplier history '
    'data model are designed in v1 to support it. Activating it requires subscribing to '
    'existing HandoffSubmitted events and deploying the agent — no v1 code changes needed.'
)
body('Tools planned:', before=40, after=20, bold=True)
bullet('analyze_photo_metadata(photo_s3_key)  →  checks EXIF timestamp vs submission timestamp; detects recycled images')
bullet('detect_duplicate_image(photo_s3_key, participant_id)  →  perceptual hash comparison against prior submissions')
bullet('score_supplier_reliability(participant_id)  →  rolling score based on on-time submission rate, deviation history, verification pass rate')
bullet('flag_for_human_review(lot_id, event_id, evidence)  →  creates a review task in the retailer dashboard')
bullet('Amazon Rekognition integration  →  image quality and content validation (v2+)')

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — DATA ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
section_h1('7', 'Data Architecture')

h2('7.1  Primary Database — Amazon RDS PostgreSQL')
aws_badge('RDS PostgreSQL 16', 'Multi-AZ deployment  |  db.t4g.medium (v1)  →  db.r8g.large (v2)')
aws_badge('RDS Proxy', 'connection pooling — prevents Lambda connection storms')
aws_badge('KMS', 'encryption at rest with AWS-managed key')

body('Core entity schema:', before=60, after=20)
mk_table(
    ['Table', 'Key Columns', 'Notes'],
    [
        ['retailers',         'id, name, brand_config (JSONB), subscription_tier',                  'brand_config: logo_url, primary_colour, display_name'],
        ['stores',            'id, retailer_id, name, address, gps_point',                          'PostGIS extension for geo queries in v2'],
        ['users',             'id, role, phone_hash, email, cognito_sub, retailer_id, store_id',    'phone stored as bcrypt hash — never plaintext'],
        ['products',          'id, retailer_id, name, category, temp_threshold_f, description',     'temp_threshold used by Alert Agent'],
        ['lots',              'id, product_id, lot_code, status, created_at, qr_published_at',      'lot_code: RL-YYYY-NNNN (retailer-canonical)'],
        ['lot_participants',  'id, lot_id, participant_id, sequence, status, invited_at, logged_at','sequence determines handoff order'],
        ['handoff_events',    'id, lot_id, participant_id, event_type, temp_f, photo_s3_key, gps, cert_number, submitted_at, hash',  'hash = SHA-256 of canonical JSON'],
        ['alert_log',         'id, lot_id, event_id, type, severity, resolved_at, resolved_by',    'All alerts, manual resolutions tracked'],
        ['notification_log',  'id, recipient_id, lot_id, type, channel, sent_at, status',           'Deduplication source for Notification Agent'],
        ['export_log',        'id, lot_id, requested_by, status, s3_url, expires_at',               'FSMA export history, presigned URL tracking'],
    ]
)

h2('7.2  Audit Log — Amazon DynamoDB')
aws_badge('DynamoDB', 'On-demand billing  |  append-only audit log table  |  no TTL')
body(
    'Every handoff event is written to DynamoDB immediately after being written to RDS. '
    'The DynamoDB table is append-only — no updates or deletes are permitted by IAM policy. '
    'This is the tamper-evident layer. If a record in RDS is modified, the hash in '
    'DynamoDB will not match — detectable by any re-verification process.',
    before=60
)
code_block(
'''Table: rootline_audit_log
  Partition key: lot_id (String)
  Sort key:      submitted_at#event_id (String)  — range queries by time

  Attributes:
    event_id      String    — matches handoff_events.id in RDS
    event_type    String    — HARVEST | COLD_STORAGE | TRANSIT | DELIVERY
    participant_id String
    submitted_at  String    — ISO8601 UTC, server-assigned
    hash          String    — SHA-256(canonical_event_json)
    verified      Boolean   — set by Verification Agent post-check
    verification_status String  — VERIFIED | ANOMALY_FLAGGED | PENDING
'''
)

h2('7.3  Object Storage — Amazon S3')
aws_badge('S3', 'three buckets  |  all encrypted with SSE-S3  |  versioning enabled')
mk_table(
    ['Bucket', 'Contents', 'Access', 'Retention'],
    [
        ['rootline-media-{acct}',   'Handoff photos, retailer logos',    'Private — Lambda only via IAM role',   'Indefinite'],
        ['rootline-exports-{acct}', 'FSMA 204 PDF + CSV exports',        'Presigned URL (7-day expiry)',          'Indefinite (audit)'],
        ['rootline-qr-{acct}',      'Consumer QR static HTML pages',     'Public read via CloudFront OAI',       'Lot lifetime'],
        ['rootline-brand-{acct}',   'Retailer logo uploads for QR pages','Public read via CloudFront OAI',       'Indefinite'],
    ]
)

h2('7.4  Cache — Amazon ElastiCache (Redis)')
aws_badge('ElastiCache Redis 7', 'single-node (v1)  →  cluster mode (v2)  |  t4g.micro')
mk_table(
    ['Cache key pattern', 'TTL', 'Purpose'],
    [
        ['session:{jwt_jti}',          '24h',   'JWT revocation list — invalidated on logout'],
        ['otp:{phone_hash}',           '10min', 'Active OTP code for supplier auth'],
        ['otp_rate:{phone_hash}',      '15min', 'Rate limit counter — max 3 OTP requests per window'],
        ['brand:{retailer_id}',        '1h',    'Retailer brand config — invalidated on PATCH'],
        ['lot_status:{lot_id}',        '5min',  'Dashboard lot status — invalidated on HandoffSubmitted'],
        ['consumer_page:{lot_id}',     '1h',    'Consumer QR page data — invalidated on lot update'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — SECURITY ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
section_h1('8', 'Security Architecture')

h2('8.1  Network Security')
aws_badge('VPC', 'all data-tier resources in private subnets — no public IP')
aws_badge('WAF', 'rate limiting (1000 req/5min per IP), SQLi, XSS rules on API Gateway')
aws_badge('Shield Standard', 'DDoS protection — included with CloudFront and API Gateway')
body(
    'Lambda functions in the API/service layer run in a VPC with private subnets. '
    'RDS, DynamoDB VPC Endpoint, ElastiCache, and Secrets Manager are accessible only '
    'from within the VPC. No data-tier resource has a public endpoint.',
    before=60
)

h2('8.2  Identity & Access')
aws_badge('Cognito', 'retailer user pool — email + password + optional MFA')
aws_badge('IAM', 'Lambda execution roles — least-privilege, per-service roles')
aws_badge('Secrets Manager', 'DB credentials, Pinpoint API key, USDA API key — auto-rotated')
bullet('Each Lambda function has its own IAM execution role — no shared credentials')
bullet('Bedrock agents have their own IAM role — can only invoke their defined tools')
bullet('S3 buckets have explicit DENY for all public access except rootline-qr and rootline-brand (served via CloudFront OAI)')
bullet('DynamoDB audit table: IAM policy denies UpdateItem and DeleteItem for all roles including admin')

h2('8.3  Data Privacy')
bullet('Phone numbers stored as bcrypt hash (cost factor 12) — never stored in plaintext')
bullet('RDS encrypted at rest with AWS KMS (aws/rds managed key)')
bullet('S3 buckets encrypted with SSE-S3')
bullet('All data in transit encrypted via TLS 1.2+ enforced at API Gateway and CloudFront')
bullet('Consumer QR page contains: farm name, region, cert type, cold chain status — no PII')
bullet('GDPR right-to-erasure: supplier phone hash deleted from users table on request; DynamoDB audit records are anonymised (participant_id nulled) — hash remains for integrity')

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — INFRASTRUCTURE & DEPLOYMENT
# ══════════════════════════════════════════════════════════════════════════════
section_h1('9', 'Infrastructure & Deployment')

h2('9.1  Infrastructure as Code')
aws_badge('AWS CDK (TypeScript)', 'all infrastructure defined as code — version controlled')
body(
    'Every AWS resource is defined in CDK stacks. Infrastructure is deployed from code, '
    'not the console. Stacks are organised by layer: NetworkStack, DataStack, ApiStack, '
    'AgentStack, FrontendStack. Each stack has explicit dependencies.',
    before=60
)
bullet('NetworkStack: VPC, subnets, security groups, VPC endpoints')
bullet('DataStack: RDS, DynamoDB, ElastiCache, S3 buckets, KMS keys')
bullet('ApiStack: API Gateway (REST + WebSocket), Lambda functions, Lambda Layers, IAM roles')
bullet('AgentStack: Bedrock agent definitions, Step Functions state machines, EventBridge rules, SQS queues')
bullet('FrontendStack: App Runner services, CloudFront distributions, Route 53 records, WAF rules')

h2('9.2  CI/CD Pipeline')
aws_badge('CodePipeline', 'three-stage pipeline: Source → Build → Deploy')
aws_badge('CodeBuild', 'runs tests, CDK synth, Next.js build')
aws_badge('CodePipeline', 'separate pipelines for each environment (dev / staging / prod)')
code_block(
'''GitHub push to main
    └─► CodePipeline triggers
         ├─► CodeBuild: npm test + cdk synth
         ├─► Deploy: dev environment (auto)
         ├─► Manual approval gate
         └─► Deploy: production environment
'''
)

h2('9.3  Environments')
mk_table(
    ['Environment', 'AWS Account', 'Purpose', 'Data'],
    [
        ['dev',     'Same account, dev prefix',  'Active development, quick iteration',          'Synthetic data only'],
        ['staging', 'Same account, stg prefix',  'Pre-release validation, load testing',         'Anonymised production-like data'],
        ['prod',    'Same account, prod prefix',  'Live customer traffic',                        'Real customer data — full encryption'],
    ]
)

h2('9.4  Observability')
aws_badge('CloudWatch', 'structured JSON logs from all Lambda functions')
aws_badge('X-Ray', 'distributed tracing — end-to-end request visibility across Lambda + Bedrock')
aws_badge('CloudWatch Alarms', 'P99 latency, error rate, agent failure rate — pages on-call')
bullet('Every Lambda function emits structured JSON logs: {request_id, lot_id, event_type, duration_ms, status}')
bullet('Bedrock agent invocations traced via X-Ray — model latency, token usage, tool call success/failure')
bullet('Key SLO alarms: API Gateway P99 > 3s, Lambda error rate > 1%, Pinpoint SMS delivery rate < 95%')
bullet('Business metrics dashboard: lots created/day, handoffs logged/day, chains completed, QR scans/day')

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — SCALABILITY ROADMAP
# ══════════════════════════════════════════════════════════════════════════════
section_h1('10', 'Scalability Roadmap — v1 → v2 → v3')
body(
    'The event-driven agent architecture means almost every v2 and v3 capability is '
    'additive — a new agent subscribes to existing events, or a new tool is added to '
    'an existing agent. No v1 data migration is required for any planned v2 feature.'
)

mk_table(
    ['Capability', 'v1', 'v2', 'v3', 'Architecture impact'],
    [
        ['Temperature logging',   'Manual entry / binary confirm', 'IoT sensor POST to /handoffs API (same endpoint)', 'AWS IoT Core → EventBridge', 'None — same event schema'],
        ['Photo verification',    'Stored, metadata checked',      'Fraud Detection Agent (Rekognition + Bedrock)',      'Advanced CV model',          'New agent, existing events'],
        ['Supplier analytics',    'Not available',                 'Supplier risk score dashboard',                     'Predictive risk model',      'New agent + new RDS table'],
        ['Blockchain anchoring',  'SHA-256 audit log (DynamoDB)',  'Anchor DynamoDB hashes to Ethereum / Hyperledger',  'Consortium chain',           'New agent, existing hashes'],
        ['Multi-language',        'English only',                  'Bedrock translate tool in Notification Agent',       'Full i18n',                  'New agent tool'],
        ['Consumer app',          'Mobile web (QR scan)',          'Retailer embeds API in their native app',           'Rootline consumer app',      'Existing API + new surface'],
        ['Retailer multi-org',    'Single org per account',        'Parent / franchise hierarchy',                      'Enterprise SSO (SAML)',      'Cognito user pool update'],
        ['Regulatory expansion',  'FSMA 204 (US)',                 'EU FIC / UK FSA traceability',                      'Global KDE mapping engine',  'Compliance Agent new tools'],
    ]
)

note(
    'v2 activation trigger: 50+ active retailer accounts or first enterprise retailer (100+ stores). '
    'v3 activation trigger: international expansion or Series A close.'
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — AWS SERVICE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
section_h1('11', 'AWS Service Summary')
mk_table(
    ['AWS Service', 'Role in Rootline v1', 'v1 Sizing'],
    [
        ['API Gateway (REST)',        'All client-server REST communication',                   'Regional, pay-per-request'],
        ['API Gateway (WebSocket)',   'Real-time dashboard updates',                            'Pay-per-message'],
        ['Lambda',                   '8 microservices + authorizers',                          'Node.js 22, 512MB–1GB, 30s timeout'],
        ['App Runner',               'Next.js hosting — supplier portal + retailer dashboard', '1 vCPU / 2GB, auto-scale 1–10'],
        ['CloudFront',               'CDN for all three web surfaces',                         'Global edge, OAI for S3'],
        ['S3',                       'Photos, exports, QR pages, brand assets',                '4 buckets, SSE-S3 encryption'],
        ['RDS PostgreSQL',           'Primary relational database',                            'db.t4g.medium, Multi-AZ, RDS Proxy'],
        ['DynamoDB',                 'Append-only tamper-evident audit log',                   'On-demand, no TTL'],
        ['ElastiCache Redis',        'Session cache, OTP, rate limits, lot status',            'cache.t4g.micro single node'],
        ['EventBridge',              'Domain event bus — agent trigger routing',               'Default bus + custom rootline bus'],
        ['SQS',                      'Agent input queues + DLQs',                              'Standard queue, 14-day retention'],
        ['Step Functions',           'Compliance Agent workflow orchestration',                 'Standard workflows'],
        ['Bedrock (Claude Haiku)',    'Verification + Alert + Notification agents',             'On-demand, us-east-1'],
        ['Bedrock (Claude Sonnet)',   'Compliance Agent (reasoning-heavy)',                     'On-demand, us-east-1'],
        ['Cognito',                  'Retailer authentication and user management',            '1 User Pool'],
        ['Pinpoint',                 'SMS — OTP, invitations, alerts, notifications',          'Pay-per-message, US long code'],
        ['KMS',                      'Encryption keys for RDS, S3',                           'AWS-managed keys (no cost)'],
        ['Secrets Manager',          'DB credentials, API keys — auto-rotation',              'Pay-per-secret/month'],
        ['WAF',                      'API protection — rate limiting, SQLi, XSS',             'WebACL on API Gateway + CloudFront'],
        ['Route 53',                 'DNS for all Rootline domains',                           'Hosted zone + A records'],
        ['Textract',                 'OCR for cert number extraction from photos',             'Pay-per-page, async API'],
        ['CloudWatch',               'Logs, metrics, alarms, dashboards',                      'Log retention: 90 days'],
        ['X-Ray',                    'Distributed tracing across Lambda + Bedrock',            'Sampling rate: 5%'],
        ['CodePipeline + CodeBuild', 'CI/CD for all services and infrastructure',             '3 pipelines (dev/stg/prod)'],
        ['CDK (TypeScript)',         'Infrastructure as code',                                 'All stacks version-controlled'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12 — OPEN ENGINEERING DECISIONS
# ══════════════════════════════════════════════════════════════════════════════
section_h1('12', 'Open Engineering Decisions')
body(
    'The following decisions must be made before engineering begins. '
    'They are documented here as open questions, not deferrals.'
)

decisions = [
    ('OD-01', 'PDF Generation Approach',
     'Two options: (a) headless Chromium via Lambda (Lambda layer, ~50MB, consistent rendering) or '
     '(b) a PDF library (pdfkit/reportlab). Chromium gives pixel-perfect branded output matching the '
     'consumer QR page design; pdfkit is lighter but requires manual layout code. '
     'Recommendation: headless Chromium for v1 — brand consistency outweighs the Lambda cold-start penalty.'),
    ('OD-02', 'Supplier Portal Hosting: App Runner vs Vercel',
     'App Runner keeps everything in AWS and simplifies IAM. Vercel offers better Next.js DX, '
     'faster Edge deployments, and built-in preview URLs for each branch. '
     'Recommendation: App Runner for v1 to keep infrastructure simple; '
     'evaluate Vercel if frontend iteration speed becomes a bottleneck.'),
    ('OD-03', 'Lot Code Format',
     'Current format: RL-YYYY-NNNN (e.g., RL-2026-0001). Needs a decision on '
     'whether to make lot codes globally unique (auto-increment across all retailers) '
     'or scoped per retailer (RL-{retailer_prefix}-YYYY-NNNN). '
     'Recommendation: globally unique — simplifies URL routing and FSMA export identity.'),
    ('OD-04', 'USDA API Rate Limits',
     'The USDA Organic Integrity Database API has undocumented rate limits. '
     'The Verification Agent must implement exponential backoff and a local cache '
     'of recently verified certificate numbers (Redis, 24h TTL) to avoid hitting limits '
     'during high-volume onboarding periods.'),
    ('OD-05', 'WebSocket Connection Management',
     'API Gateway WebSocket connections time out after 10 minutes of inactivity. '
     'The retailer dashboard must implement a heartbeat (ping every 9 minutes) and '
     'reconnect logic. Decision needed: implement in client JavaScript or use '
     'a managed solution like AWS AppSync (adds GraphQL subscription model).'),
    ('OD-06', 'Multi-Region Strategy',
     'v1 launches in us-east-1. If a UK or EU retailer onboards, data residency requirements '
     '(GDPR) may require a eu-west-1 deployment. Decision needed before signing any non-US customer: '
     'single-region with data processing agreements, or multi-region from day one.'),
]

for code, title, detail in decisions:
    h3(f'{code} — {title}')
    body(detail, before=20, after=80)

# ── Footer ─────────────────────────────────────────────────────────────────────
thin_rule(before=120, after=60)
fp = doc.add_paragraph(); sp(fp, before=0, after=0)
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run('Rootline  ·  High-Level Design v1.0  ·  June 2026  ·  Confidential  ·  AWS Account 741705262274')
r.font.name='Arial'; r.font.size=Pt(8); r.font.italic=True; r.font.color.rgb=GREY

# ── Save ───────────────────────────────────────────────────────────────────────
out = Path('/Users/avina/Documents/rootline-app/docs/engineering/Rootline_HLD_v1.docx')
doc.save(out)
print(f'Saved: {out}')
